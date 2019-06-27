#!/usr/bin/env python
# coding: utf-8

# # Resume Creator
# ## User Requirements: A plaintext resume in JSON format, see example:
# #### https://github.com/davidschneider04/resume/raw/master/plaintext_resume.json

# #### Imports

# In[1]:

def main():
    import docx, math, json, re
    import urllib.request


    # #### Load a JSON with plaintext resume data. Theoretically, this could be done from the local machine.

    # In[2]:


    #ptexturl = 'https://github.com/davidschneider04/resume/raw/master/plaintext_resume.json'
    ptexturl = "/Users/kutch/resume/plaintext_resume.json"
    ptext = json.load(urllib.request.urlopen(ptexturl))


    # ## Given the information in the plaintext resume, populate a Word doc with proper styling

    # #### Custom Styles Used:
    # <ul>
    #     <li>ResumeSectionHeader</li>
    #     <li>ResumeName</li>
    #     <li>ResumeContactLine</li>
    #     <li>ResumeSummary</li>
    #     <li>ResumeCompanyHeader</li>
    #     <li>ResumeCompanyDescription</li>
    #     <li>ResumeAccomplishment2</li>
    #     <li>ResumePositionDescription</li>
    #     <li>ResumeWordJumble</li>
    # </ul>

    # #### Create the Word doc

    # In[3]:


    #this document holds all the styles
    doc = docx.Document('BlankTemplate.docx')

    def create_section_header(name):
        section_header_style = doc.styles['ResumeSectionHeader']
        section_header_text = name
        section_header_p = doc.add_paragraph(section_header_text)
        section_header_p.style = section_header_style
        
    def delete_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None


    # #### Name Info

    # In[4]:


    name_text = ptext['name']
    name_style = doc.styles['ResumeName']
    name_p = doc.add_paragraph(name_text)
    name_p.style = name_style


    # #### Delete the blank paragraph added by default

    # In[5]:


    delete_paragraph(doc.paragraphs[0])


    # #### Contact Info

    # In[6]:


    contact_text = ''
    for entry, val in ptext['contact'].items():
        contact_text += val
        contact_text += '\t' if entry != 'website' else ''
    contact_style = doc.styles['ResumeContactLine']
    contact_p = doc.add_paragraph(contact_text)
    contact_p.style = contact_style


    # #### Summary Info

    # In[7]:


    create_section_header("Summary")
    summary_text = ptext['summary']
    summary_style = doc.styles['ResumeSummary']
    summary_p = doc.add_paragraph(summary_text)
    summary_p.style = summary_style


    # #### Work Experience Info

    # In[8]:


    create_section_header("Professional Experience")
    jobtitle_style = doc.styles['ResumeCompanyHeader']
    jobdescription_style = doc.styles['ResumeCompanyDescription']
    jobaccomplishment_style = doc.styles['ResumeAccomplishment2']
    for job in ptext['experience']:
        jobheader_p = doc.add_paragraph('')
        jobheader_p.add_run(job['title']).bold = True
        jobheader_p.add_run(" \u2014 " + job['company'] + '\t' 
                            + job['start'] + " \u2014 " + job['end'])
        jobheader_p.style = jobtitle_style
        jobdescription_text = job['company_description']
        jobdescription_p = doc.add_paragraph(jobdescription_text)
        jobdescription_p.style = jobdescription_style
        for accomplishment in job['achievements']:
            jobaccomplishment_text = accomplishment
            jobaccomplishment_p = doc.add_paragraph(jobaccomplishment_text)
            jobaccomplishment_p.style = jobaccomplishment_style


    # #### Education / Certification Info

    # In[9]:


    create_section_header("Education / Certifications")
    education_style = doc.styles['ResumeCompanyHeader']
    for degree in ptext['education']:
        education_p = doc.add_paragraph('')
        education_p.add_run(degree['name']).bold = True
        education_p.add_run('\t' + degree['earndate'])
        education_p.style = education_style
        if degree['description']:
            #kind of extra
            collegedegree = re.compile('[BM]\.?[SA]\.?\s.+(?=,)')
            if re.search(collegedegree, degree['description']) and re.search('minor', degree['description'], flags=re.IGNORECASE):
                educationdescription_p = doc.add_paragraph('')
                collegedegreetext = re.findall(collegedegree, degree['description'])[0]
                educationdescription_p.add_run(collegedegreetext).bold = True
                othertext = degree['description'][degree['description'].index(','):]
                educationdescription_p.add_run(othertext)
            else:
                educationdescription_text = degree['description']
                educationdescription_p = doc.add_paragraph(educationdescription_text)
            educationdescription_p.style = doc.styles['ResumePositionDescription']


    # #### Skills Info

    # In[10]:


    create_section_header("Skills")
    skills_style = doc.styles["ResumeWordJumble"]
    tblcols = 4
    tblrows = math.ceil(len(ptext['skills'])/tblcols)
    skilltable = doc.add_table(rows=tblrows, cols=tblcols)
    skilltable.style = doc.styles['ResumeSkillTable']
    #skilltable.style.paragraph_format.left_indent = docx.shared.Inches(0.38)
    i, j = 0, 0
    for skill in ptext['skills']:
        skilltable.rows[i].cells[j].text = skill
        skilltable.rows[i].cells[j].paragraphs[0].style = skills_style
        j+=1
        if j%(tblcols) == 0:
            j=0
            i+=1


    # #### Save out doc

    # In[11]:


    doc.save('DavidSchneiderResume.docx')


if __name__ == "__main__":
    # execute only if run as a script
    main()
